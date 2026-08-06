"""One-message deterministic document processor for an Azure Container Apps Job."""

from __future__ import annotations

import csv
import hashlib
import hmac
import io
import json
import mimetypes
import os
import re
import subprocess
import tempfile
import time
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any

import requests
from azure.identity import DefaultAzureCredential
from azure.servicebus import ServiceBusClient
from azure.core.exceptions import ResourceExistsError
from azure.storage.blob import BlobServiceClient, ContentSettings
from docx import Document
from openpyxl import load_workbook
from PIL import Image, ImageChops, ImageDraw, ImageOps, ImageSequence

PROCESSOR_VERSION = "tiecamel-document-processor/2"


def required(name: str) -> str:
    value = os.getenv(name)
    if not value:
        raise RuntimeError(f"{name} is required")
    return value


def sha256(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def normalize(text: str) -> str:
    return "\n".join(line.rstrip() for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n")).strip()


def blob_parts(reference: str) -> tuple[str, str]:
    value = reference.removeprefix("azure://")
    container, path = value.split("/", 1)
    return container, path


@dataclass
class Artifact:
    kind: str
    path: Path
    mime: str
    metadata: dict[str, Any]


class Processor:
    def __init__(self) -> None:
        self.credential = DefaultAzureCredential()
        self.blobs = BlobServiceClient(required("AZURE_STORAGE_BLOB_URL"), credential=self.credential)

    def download(self, reference: str) -> bytes:
        container, path = blob_parts(reference)
        return self.blobs.get_blob_client(container, path).download_blob().readall()

    def upload_artifact(self, command: dict[str, Any], artifact: Artifact) -> dict[str, Any]:
        data = artifact.path.read_bytes()
        digest = sha256(data)
        path = f"document-artifacts/{command['organizationId']}/{command['repositoryId']}/{command['uploadSessionId']}/{digest}/{artifact.path.name}"
        client = self.blobs.get_blob_client("processed", path)
        try:
            client.upload_blob(data, overwrite=False, content_settings=ContentSettings(content_type=artifact.mime))
        except ResourceExistsError:
            existing = client.download_blob().readall()
            if sha256(existing) != digest:
                raise RuntimeError("Existing content-addressed artifact has unexpected bytes")
        return {
            "kind": artifact.kind,
            "objectRef": f"azure://processed/{path}",
            "sha256": digest,
            "processorVersion": PROCESSOR_VERSION,
            "metadata": artifact.metadata,
        }

    def process(self, command: dict[str, Any]) -> dict[str, Any]:
        proposed = self.download(command["azureBlobRef"])
        if len(proposed) != command["expectedSize"] or sha256(proposed) != command["expectedSha256"].lower():
            raise RuntimeError("Uploaded bytes do not match the authorized size and checksum")
        base = self.download(command["baseVersionRef"]) if command.get("baseVersionRef") else None
        with tempfile.TemporaryDirectory(prefix="tiecamel-") as work:
            root = Path(work)
            proposed_path = root / safe_name(command["fileName"])
            proposed_path.write_bytes(proposed)
            detected_mime = detect_mime(proposed_path)
            if detected_mime != command["mimeType"]:
                raise RuntimeError(f"MIME mismatch: declared {command['mimeType']}, detected {detected_mime}")
            scan(proposed_path)
            proposed_doc = extract(proposed_path, detected_mime, root / "proposed")
            base_doc = None
            if base is not None:
                base_path = root / f"base{proposed_path.suffix}"
                base_path.write_bytes(base)
                scan(base_path)
                if detect_mime(base_path) != detected_mime:
                    raise RuntimeError("Accepted and proposed documents use different formats")
                base_doc = extract(base_path, detected_mime, root / "base")
            text_diff = line_diff(base_doc["text"] if base_doc else "", proposed_doc["text"])
            structured = structured_diff(base_doc, proposed_doc)
            artifacts = proposed_doc["artifacts"]
            if base_doc:
                artifacts += base_doc["artifacts"]
                artifacts += visual_overlays(base_doc, proposed_doc, root / "overlays")
            artifact_results = [self.upload_artifact(command, artifact) for artifact in artifacts]
            normalized = normalize(proposed_doc["text"])
            base_normalized = normalize(base_doc["text"]) if base_doc else None
            return {
                "sha256": sha256(proposed),
                "normalizedSha256": sha256(normalized.encode()),
                "detectedMimeType": detected_mime,
                "malwareScan": "clean",
                "extracted": proposed_doc.get("fields", []),
                "findings": structured,
                "textDiff": text_diff,
                "visualManifestKey": next((item["objectRef"] for item in artifact_results if item["kind"] == "page-render"), None),
                "artifacts": artifact_results,
                "diff": {
                    "format": "tiecamel-document-diff/v2",
                    "baseContentSha256": sha256(base) if base else None,
                    "proposedContentSha256": sha256(proposed),
                    "baseNormalizedSha256": sha256(base_normalized.encode()) if base_normalized is not None else None,
                    "proposedNormalizedSha256": sha256(normalized.encode()),
                    "stats": {
                        "additions": sum(1 for item in text_diff if item["type"] == "added"),
                        "deletions": sum(1 for item in text_diff if item["type"] == "removed"),
                        "changes": len(structured),
                    },
                },
                "processorVersion": PROCESSOR_VERSION,
            }


def scan(path: Path) -> None:
    qpdf = subprocess.run(["qpdf", "--check", str(path)], capture_output=True, text=True) if path.suffix.lower() == ".pdf" else None
    if qpdf and qpdf.returncode not in (0, 3):
        raise RuntimeError(f"Malformed or encrypted PDF: {qpdf.stderr.strip()}")
    result = subprocess.run(["clamscan", "--no-summary", str(path)], capture_output=True, text=True)
    if result.returncode == 1:
        raise RuntimeError("Malware detected")
    if result.returncode > 1 and os.getenv("REQUIRE_MALWARE_SCANNER", "true") == "true":
        raise RuntimeError(f"Malware scanner failed: {result.stderr.strip()}")


def detect_mime(path: Path) -> str:
    head = path.read_bytes()[:8192]
    if head.startswith(b"%PDF-"):
        return "application/pdf"
    if head.startswith(b"\x89PNG\r\n\x1a\n"):
        return "image/png"
    if head.startswith(b"\xff\xd8\xff"):
        return "image/jpeg"
    if head.startswith((b"II*\x00", b"MM\x00*")):
        return "image/tiff"
    if head.startswith(b"PK\x03\x04"):
        with zipfile.ZipFile(path) as archive:
            names = set(archive.namelist())
            if "word/document.xml" in names:
                return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            if "xl/workbook.xml" in names:
                return "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        raise RuntimeError("Unsupported or malformed Office archive")
    if b"\x00" not in head:
        try:
            head.decode("utf-8")
            return "text/csv"
        except UnicodeDecodeError:
            pass
    raise RuntimeError("Unsupported or unrecognized document format")


def extract(path: Path, mime: str, output: Path) -> dict[str, Any]:
    output.mkdir(parents=True, exist_ok=True)
    if mime == "application/pdf":
        return extract_pdf(path, output)
    if mime.endswith("wordprocessingml.document"):
        return extract_docx(path, output)
    if mime.endswith("spreadsheetml.sheet"):
        return extract_xlsx(path, output)
    if mime == "text/csv":
        return extract_csv(path, output)
    if mime.startswith("image/"):
        return extract_image(path, output)
    raise RuntimeError(f"Unsupported document type {mime}")


def extract_pdf(path: Path, output: Path) -> dict[str, Any]:
    text_path = output / "content.txt"
    subprocess.run(["pdftotext", "-layout", str(path), str(text_path)], check=True)
    render_prefix = output / "page"
    subprocess.run(["pdftoppm", "-png", "-r", "144", str(path), str(render_prefix)], check=True)
    text = text_path.read_text(errors="replace")
    pages = sorted(output.glob("page-*.png"))
    if not text.strip():
        text = "\n\f\n".join(ocr(page) for page in pages)
        text_path.write_text(text)
    artifacts = [Artifact("text", text_path, "text/plain", {"source": "pdf"})]
    artifacts += [Artifact("page-render", page, "image/png", {"page": index + 1, "side": output.name}) for index, page in enumerate(pages)]
    return {"text": text, "artifacts": artifacts, "pages": pages, "cells": {}, "fields": extract_labeled_fields(text)}


def extract_docx(path: Path, output: Path) -> dict[str, Any]:
    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table_index, table in enumerate(document.tables, 1):
        parts.append(f"[Table {table_index}]")
        parts.extend("\t".join(cell.text for cell in row.cells) for row in table.rows)
    for section_index, section in enumerate(document.sections, 1):
        parts.append(f"[Header {section_index}]")
        parts.extend(paragraph.text for paragraph in section.header.paragraphs)
        parts.append(f"[Footer {section_index}]")
        parts.extend(paragraph.text for paragraph in section.footer.paragraphs)
    text_path = output / "content.txt"
    text_path.write_text("\n".join(parts))
    subprocess.run(["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", str(output), str(path)], check=True)
    pdf = output / f"{path.stem}.pdf"
    rendered = extract_pdf(pdf, output / "rendered") if pdf.exists() else {"artifacts": [], "pages": []}
    text = text_path.read_text()
    return {"text": text, "artifacts": [Artifact("text", text_path, "text/plain", {"source": "docx"})] + rendered["artifacts"], "pages": rendered["pages"], "cells": {}, "fields": extract_labeled_fields(text)}


def extract_xlsx(path: Path, output: Path) -> dict[str, Any]:
    workbook = load_workbook(path, read_only=False, data_only=False, keep_links=False)
    cells: dict[str, dict[str, Any]] = {}
    lines: list[str] = []
    for sheet in workbook.worksheets:
        lines.append(f"[Sheet:{sheet.title};hidden={sheet.sheet_state != 'visible'}]")
        for table_name, table in sorted(sheet.tables.items()):
            lines.append(f"[Table:{sheet.title}!{table_name};ref={getattr(table, 'ref', table)}]")
        for key, dimension in sorted(sheet.row_dimensions.items()):
            if dimension.hidden:
                lines.append(f"[HiddenRow:{sheet.title}!{key}]")
        for key, dimension in sorted(sheet.column_dimensions.items()):
            if dimension.hidden:
                lines.append(f"[HiddenColumn:{sheet.title}!{key}]")
        for row in sheet.iter_rows():
            for cell in row:
                if cell.value is None:
                    continue
                key = f"{sheet.title}!{cell.coordinate}"
                value = cell.value
                cells[key] = {"value": str(value), "type": cell.data_type, "formula": str(value) if cell.data_type == "f" else None}
                lines.append(f"{key}\t{cell.data_type}\t{value}")
    for defined_name in sorted(workbook.defined_names.values(), key=lambda item: item.name):
        lines.append(f"[NamedRange:{defined_name.name};value={defined_name.attr_text}]")
    text_path = output / "workbook.tsv"
    text_path.write_text("\n".join(lines))
    return {"text": text_path.read_text(), "artifacts": [Artifact("table", text_path, "text/tab-separated-values", {"sheets": workbook.sheetnames})], "pages": [], "cells": cells}


def extract_csv(path: Path, output: Path) -> dict[str, Any]:
    raw = path.read_text(errors="replace")
    dialect = csv.Sniffer().sniff(raw[:8192], delimiters=",\t;|")
    rows = list(csv.reader(io.StringIO(raw), dialect))
    normalized_path = output / "table.csv"
    with normalized_path.open("w", newline="") as handle:
        csv.writer(handle, lineterminator="\n").writerows(rows)
    cells = {f"row-{row_index + 1}:{column_index + 1}": {"value": value, "type": infer_type(value)} for row_index, row in enumerate(rows) for column_index, value in enumerate(row)}
    return {"text": normalized_path.read_text(), "artifacts": [Artifact("table", normalized_path, "text/csv", {"delimiter": dialect.delimiter, "rows": len(rows), "rowMatching": "position-fallback"})], "pages": [], "cells": cells}


def extract_image(path: Path, output: Path) -> dict[str, Any]:
    source = Image.open(path)
    pages = []
    artifacts = []
    text_parts = []
    for index, frame in enumerate(ImageSequence.Iterator(source), 1):
        image = ImageOps.exif_transpose(frame).convert("RGB")
        normalized_path = output / f"image-{index}.png"
        image.save(normalized_path)
        pages.append(normalized_path)
        text_parts.append(ocr(normalized_path))
        artifacts.append(Artifact("page-render", normalized_path, "image/png", {"page": index, "side": output.name, "width": image.width, "height": image.height}))
    text = "\n\f\n".join(text_parts)
    text_path = output / "ocr.txt"
    text_path.write_text(text)
    artifacts.append(Artifact("text", text_path, "text/plain", {"source": "ocr"}))
    return {"text": text, "artifacts": artifacts, "pages": pages, "cells": {}, "fields": extract_labeled_fields(text)}


def ocr(path: Path) -> str:
    return subprocess.run(["tesseract", str(path), "stdout"], check=True, capture_output=True, text=True).stdout


def extract_labeled_fields(text: str) -> list[dict[str, Any]]:
    """Extract deterministic compliance fields while retaining page provenance."""
    patterns = [
        (
            re.compile(r"PROPERTY\s+ACCOUNT\s+NOTICE\s+DATE\s+([0-9-]+)\s+([A-Za-z]+\s+\d{1,2},\s+\d{4})", re.IGNORECASE),
            (("property-account", 1), ("notice-date", 2)),
        ),
        (
            re.compile(r"AMOUNT\s+DUE\s+PENALTY(?:\s+AND)?\s+INTEREST\s+(\$[\d,]+(?:\.\d{2})?)\s+(\$[\d,]+(?:\.\d{2})?)", re.IGNORECASE),
            (("amount-due", 1), ("penalty-and-interest", 2)),
        ),
        (
            re.compile(r"RESPONSE\s+DEADLINE\s+EXEMPTION\s+STATUS\s+([A-Za-z]+\s+\d{1,2},\s+\d{4})\s+([^\n\r]+)", re.IGNORECASE),
            (("response-deadline", 1), ("exemption-status", 2)),
        ),
    ]
    fields: list[dict[str, Any]] = []
    for page_number, page in enumerate(text.split("\f"), 1):
        for pattern, captures in patterns:
            match = pattern.search(page)
            if not match:
                continue
            for field, group in captures:
                fields.append({
                    "field": field,
                    "value": " ".join(match.group(group).split()),
                    "provenance": f"page {page_number}",
                    "confidence": 1.0,
                })
    return fields


def structured_diff(base: dict[str, Any] | None, proposed: dict[str, Any]) -> list[dict[str, Any]]:
    if not base:
        return []
    findings: list[dict[str, Any]] = []
    keys = sorted(set(base.get("cells", {})) | set(proposed.get("cells", {})))
    for key in keys:
        before = base.get("cells", {}).get(key)
        after = proposed.get("cells", {}).get(key)
        if before != after:
            findings.append({"field": key, "before": json.dumps(before, sort_keys=True) if before else None, "after": json.dumps(after, sort_keys=True) if after else None, "provenance": key, "severity": "warning", "source": "deterministic"})
    base_fields = {item["field"]: item for item in base.get("fields", [])}
    proposed_fields = {item["field"]: item for item in proposed.get("fields", [])}
    for key in sorted(set(base_fields) | set(proposed_fields)):
        before = base_fields.get(key)
        after = proposed_fields.get(key)
        if (before or {}).get("value") == (after or {}).get("value"):
            continue
        findings.append({
            "field": key,
            "before": before.get("value") if before else None,
            "after": after.get("value") if after else None,
            "provenance": (after or before or {}).get("provenance", "document text"),
            "severity": "critical" if key == "response-deadline" else "warning",
            "source": "deterministic",
        })
    return findings[:10000]


def line_diff(before: str, after: str) -> list[dict[str, str]]:
    import difflib
    result = []
    for line in difflib.ndiff(normalize(before).splitlines(), normalize(after).splitlines()):
        if line.startswith("? "):
            continue
        result.append({"type": "added" if line.startswith("+ ") else "removed" if line.startswith("- ") else "unchanged", "content": line[2:]})
    return result[:20000]


def visual_overlays(base: dict[str, Any], proposed: dict[str, Any], output: Path) -> list[Artifact]:
    output.mkdir(parents=True, exist_ok=True)
    artifacts = []
    for index, (left_path, right_path) in enumerate(zip(base.get("pages", []), proposed.get("pages", [])), 1):
        left = Image.open(left_path).convert("RGB")
        right = Image.open(right_path).convert("RGB").resize(left.size)
        diff = ImageChops.difference(left, right)
        mask = diff.convert("L").point(lambda pixel: 180 if pixel > 18 else 0)
        overlay = right.copy()
        red = Image.new("RGB", right.size, (220, 38, 38))
        overlay.paste(Image.blend(right, red, 0.45), mask=mask)
        path = output / f"overlay-{index}.png"
        overlay.save(path)
        artifacts.append(Artifact("page-render", path, "image/png", {"page": index, "side": "overlay", "method": "pixel-difference"}))
    return artifacts


def infer_type(value: str) -> str:
    if re.fullmatch(r"[-+]?\d+(\.\d+)?", value.strip()):
        return "number"
    return "string"


def safe_name(value: str) -> str:
    return re.sub(r"[^A-Za-z0-9._-]", "_", value)


def callback(payload: dict[str, Any]) -> None:
    body = json.dumps(payload, separators=(",", ":"))
    timestamp = str(int(time.time() * 1000))
    signature = hmac.new(required("CONVEX_CALLBACK_SECRET").encode(), f"{timestamp}.{body}".encode(), hashlib.sha256).hexdigest()
    response = requests.post(required("CONVEX_PROCESSING_CALLBACK_URL"), data=body, headers={"content-type": "application/json", "x-tiecamel-timestamp": timestamp, "x-tiecamel-signature": signature}, timeout=30)
    response.raise_for_status()


def main() -> None:
    processor = Processor()
    with ServiceBusClient(required("AZURE_SERVICE_BUS_NAMESPACE"), credential=processor.credential) as client:
        with client.get_queue_receiver(queue_name="processing", max_wait_time=30, prefetch_count=1) as receiver:
            messages = receiver.receive_messages(max_message_count=1, max_wait_time=30)
            if not messages:
                return
            message = messages[0]
            command = json.loads(str(message))
            try:
                result = processor.process(command)
            except Exception as error:
                try:
                    callback({"uploadSessionId": command["uploadSessionId"], "idempotencyKey": command["idempotencyKey"], "succeeded": False, "error": {"code": "PROCESSING_FAILED", "message": str(error), "retryable": True}, "completedAt": time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime())})
                finally:
                    receiver.abandon_message(message)
                raise
            try:
                callback({"uploadSessionId": command["uploadSessionId"], "idempotencyKey": command["idempotencyKey"], "succeeded": True, "result": result, "completedAt": time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime())})
                receiver.complete_message(message)
            except Exception:
                # Processing succeeded. A callback transport or control-plane
                # error must be retried as the same successful result; it is
                # not a document-processing failure.
                receiver.abandon_message(message)
                raise


if __name__ == "__main__":
    main()
